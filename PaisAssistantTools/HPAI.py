# =============================================================================
# H-PAI v1.0.4 — Hypothesis, Plan, Action, Interpretation of Results
# =============================================================================
# Author: Pai Surya Darshan
# Created: 2024-04-201
#
# Purpose:
#   Edit a structured H-PAI Word template by replacing placeholders with
#   user-provided content and optionally inserting selected figure(s).
#
#   H-PAI stands for:
#   - Hypothesis
#   - Plan
#   - Action
#   - Interpretation of Results
#
# Last Updated: 2025-10-01
#
# What's New in v1.0.4:
# - Improved handling of figure insertion
# - Adapted for Master's research 
#   - Updated experiment name to "El índice de Gini"
#   - Updated Template
#   - Updated Template path to "HPAI_Template_MA.docx"
# =============================================================================

from docx import Document
from docx.shared import Inches
from datetime import datetime
from pathlib import Path
import tempfile
import matplotlib.pyplot as plt

# -----------------------------------------------------------------------------
# Function: _replace_text_everywhere
# Purpose:
#   Replace text across paragraphs and table cells in a Word document.
#
# Inputs:
#   doc          -> python-docx Document object
#   replace_map  -> dict of {old_text: new_text}
#
# Output:
#   In-place edited document
#
# Notes:
#   This is a simple run-wise replacement. Placeholders should be kept as a
#   single uninterrupted text block inside Word for best reliability.
# -----------------------------------------------------------------------------
def _replace_text_everywhere(doc, replace_map):
    def replace_in_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)

        changed = full_text
        for old, new in replace_map.items():
            changed = changed.replace(old, new)

        if changed != full_text:
            if paragraph.runs:
                paragraph.runs[0].text = changed
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(changed)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


# -----------------------------------------------------------------------------
# Function: _make_update_type_line
# Purpose:
#   Build the update-type line with the selected checkbox ticked.
#
# Inputs:
#   update_code  -> None, "WU", "BW", or "MU"
#
# Output:
#   String for the update type line
# -----------------------------------------------------------------------------
def _make_update_type_line(update_code=None):
    selected = (update_code or "None").strip().upper()

    valid = {"NONE", "WU", "BW", "MU"}
    if selected not in valid:
        selected = "NONE"

    def box(code):
        return "✅" if selected == code else "☐"

    return (
        f"Update Type: "
        f"{box('NONE')} None  "
        f"{box('WU')} WU (Weekly)  "
        f"{box('BW')} BW (Biweekly)  "
        f"{box('MU')} MU (Monthly)"
    )

# -----------------------------------------------------------------------------
# Function: _insert_figures
# Purpose:
#   Insert figure(s) exactly at the {{Figure}} placeholder location.
#
# Inputs:
#   doc         -> python-docx Document object
#   figures     -> list of figure objects and/or image paths
#   max_width   -> maximum figure width in inches
#
# Output:
#   In-place edited document
#
# Notes:
#   - Uses run.add_picture(...) so the image is inserted at the placeholder,
#     not at the end of the document.
#   - Supports matplotlib-like figures with .savefig()
#   - Supports image file paths
# -----------------------------------------------------------------------------
def _insert_figures(doc, figures, max_width=6.0):
    if not figures:
        _replace_text_everywhere(doc, {"{{Figure}}": ""})
        return

    placeholder_found = False

    for paragraph in doc.paragraphs:
        if "{{Figure}}" in paragraph.text:
            placeholder_found = True

            # Clear the placeholder from all runs
            for run in paragraph.runs:
                run.text = run.text.replace("{{Figure}}", "")

            # Insert each figure exactly here
            for i, fig in enumerate(figures, start=1):

                # Image run
                image_run = paragraph.add_run()

                if hasattr(fig, "savefig"):
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    fig.savefig(tmp.name, bbox_inches="tight", dpi=300)
                    image_run.add_picture(tmp.name, width=Inches(max_width))
                else:
                    image_run.add_picture(str(fig), width=Inches(max_width))

                # Figure label
                label_run = paragraph.add_run(f"\nFigure {i}\n")
                
                # Optional spacing after each figure
                paragraph.add_run("\n")

            break

    if not placeholder_found:
        # Fallback: append near end only if placeholder missing
        doc.add_paragraph("Figure(s):")
        for i, fig in enumerate(figures, start=1):
            p = doc.add_paragraph()
            p.add_run(f"Figure {i}\n")
            r = p.add_run()

            if hasattr(fig, "savefig"):
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                fig.savefig(tmp.name, bbox_inches="tight", dpi=300)
                r.add_picture(tmp.name, width=Inches(max_width))
            else:
                r.add_picture(str(fig), width=Inches(max_width))


# -----------------------------------------------------------------------------
# Function: fill_hpai_template
# Purpose:
#   Fill an H-PAI Word template with user content, update type, experiment name,
#   date, and optional figure(s).
#
# Inputs:
#   template_path                   -> path to .docx template
#   save_location                   -> folder to save output
#   experiment_name                 -> experiment name used in doc + filename
#   Assumption                      -> str
#   Expected_Behaviour              -> str
#   Key_Variables                   -> str
#   Methodology                     -> str
#   Data_Loading                    -> str
#   Tools                           -> str
#   Observed_Trends                 -> str
#   Deviations_from_Expectation     -> str
#   Mechanistic_Insight             -> str
#   Link_to_Hypothesis              -> str
#   figures                         -> optional list of figures / image paths
#   update_code                     -> None / "WU" / "BW" / "MU"
#   date_str                        -> optional override date
#
# Output:
#   Saved .docx file path
#
# Notes:
#   Output naming:
#       YYYY_MM_DD_experiment_name.docx
#       YYYY_MM_DD_WU_experiment_name.docx
#       YYYY_MM_DD_BW_experiment_name.docx
#       YYYY_MM_DD_MU_experiment_name.docx
# -----------------------------------------------------------------------------
def fill_hpai_template(
    save_location,
    experiment_name,
    
    Assumption="",
    Expected_Behaviour="",
    Key_Variables="",
    Methodology="",
    Data_Loading="",
    Tools="",
    Observed_Trends="",
    Deviations_from_Expectation="",
    Mechanistic_Insight="",
    Link_to_Hypothesis="",
    figures=None,
    date_str=None,

    update_code="BW",  # Default to Biweekly update

    template_path = "PaisAssistantTools/Template/HPAI_Temp.docx",
):
    doc = Document(template_path)


    date_str = date_str or datetime.now().strftime("%Y-%m-%d")
    update_line = _make_update_type_line(update_code)

    replace_map = {
        "{{YYYY-MM-DD}}": date_str,
        "{{Experiment_Name}}": experiment_name,
        "{{Assumption}}": Assumption,
        "{{Expected_Behaviour}}": Expected_Behaviour,
        "{{Key_Variables}}": Key_Variables,
        "{{Methodology}}": Methodology,
        "{{Data_Loading}}": Data_Loading,
        "{{Tools}}": Tools,
        "{{Observed_Trends}}": Observed_Trends,
        "{{Deviations_from_Expectation}}": Deviations_from_Expectation,
        "{{Mechanistic_Insight}}": Mechanistic_Insight,
        "{{Link_to_Hypothesis}}": Link_to_Hypothesis,
        "{{Update_Type}}": update_line,
    }

    _replace_text_everywhere(doc, replace_map)
    _insert_figures(doc, figures)

    date_for_filename = date_str or datetime.now().strftime("%Y_%m_%d")
    clean_name = str(experiment_name).strip().replace(" ", "_")

    selected = (update_code or "").strip().upper()
    if selected in {"WU", "BW", "MU"}:
        filename = f"{date_for_filename}_{selected}_{clean_name}.docx"
    else:
        filename = f"{date_for_filename}_{clean_name}.docx"

    save_path = Path(save_location) / filename
    save_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)

    print(f"Saved → {save_path}")
    return str(save_path)

if __name__ == "__main__":

    # # Example usage
    # fill_hpai_template(
    #     template_path="PaisAssistantTools/Template/HPAI_Temp.docx",
    #     save_location="outputs/HPAI_Docs/",
    #     experiment_name="gini_descriptor_test",
    #     update_code="BW",
    #     Assumption="Local inequality may track thermodynamic response.",
    #     Expected_Behaviour="Higher descriptor shift should align with higher reaction energy.",
    #     Key_Variables="ΔG_n, ΔG_p, Δη",
    #     Methodology="Compute descriptors and inspect PCA structure.",
    #     Data_Loading="Loaded RedDB subset and cleaned key columns.",
    #     Tools="Python, pandas, matplotlib, RDKit, Plotly",
    #     Observed_Trends="A weak directional trend is visible.",
    #     Deviations_from_Expectation="Mid-range systems remain noisy.",
    #     Mechanistic_Insight="Hardness contribution appears coupled to redistribution.",
    #     Link_to_Hypothesis="Partially supported.",
    #     figures=["PaisAssistantTools/images/test_seqa_dashboard.png", "PaisAssistantTools/images/SEQA_16.png"],
    #     date_str="2025-10-01"
    # )

    pass